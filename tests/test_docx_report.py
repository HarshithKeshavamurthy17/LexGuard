"""Test Word report generation."""

import io
from datetime import datetime
from docx import Document
from lexguard.models.contract import Contract
from lexguard.models.clause import Clause, ClauseType
from lexguard.reports.docx_report import generate_docx_report


def test_generate_docx_report():
    """Test that a valid Word document is generated."""
    # Create dummy contract
    contract = Contract(
        id="test-123",
        title="Test Contract",
        original_filename="test.pdf",
        text="This is a test contract.",
        clauses=[
            Clause(
                id="c1",
                contract_id="test-123",
                index=0,
                text="This agreement shall terminate upon notice.",
                clause_type=ClauseType.TERMINATION,
                risk_score=0.8,
                risk_level="high"
            ),
            Clause(
                id="c2",
                contract_id="test-123",
                index=1,
                text="Party A shall pay Party B $100.",
                clause_type=ClauseType.PAYMENT,
                risk_score=0.2,
                risk_level="low"
            )
        ]
    )

    risk_data = {
        "high_risk_count": 1,
        "medium_risk_count": 0,
        "low_risk_count": 1,
        "total_clauses": 2
    }

    # Generate report
    docx_bytes = generate_docx_report(contract, risk_data)

    # Verify it's a valid docx
    assert docx_bytes is not None
    assert len(docx_bytes) > 0

    # Load back with python-docx to verify content
    doc_stream = io.BytesIO(docx_bytes)
    doc = Document(doc_stream)

    # Check title
    assert doc.paragraphs[0].text == "LexGuard Contract Analysis"

    # Check content
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "Test Contract" not in full_text # Title is not in body, filename is
    assert "test.pdf" in full_text
    assert "High Risk Issues: 1" in full_text
    assert "This agreement shall terminate upon notice." in full_text
